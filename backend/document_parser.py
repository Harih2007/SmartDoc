"""
Document Parser Module
Handles parsing of PDF, TXT, and Markdown files into text chunks.
"""

import os
import re
from dataclasses import dataclass, field
from typing import List

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


@dataclass
class DocumentChunk:
    """Represents a chunk of document text with metadata."""
    text: str
    chunk_index: int
    source_file: str
    page_number: int = 0
    start_char: int = 0
    end_char: int = 0
    metadata: dict = field(default_factory=dict)


class DocumentParser:
    """Parses various document formats and splits them into searchable chunks."""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}
    CHUNK_SIZE = 500  # characters
    CHUNK_OVERLAP = 100  # characters

    @classmethod
    def parse(cls, file_path: str, file_bytes: bytes = None, filename: str = None) -> List[DocumentChunk]:
        """Parse a document file into chunks."""
        if filename is None:
            filename = os.path.basename(file_path)

        ext = os.path.splitext(filename)[1].lower()

        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {cls.SUPPORTED_EXTENSIONS}")

        if ext == ".pdf":
            raw_text, page_texts = cls._parse_pdf(file_path, file_bytes)
        elif ext == ".docx":
            raw_text, page_texts = cls._parse_docx(file_path, file_bytes)
        else:
            raw_text, page_texts = cls._parse_text(file_path, file_bytes)

        chunks = cls._chunk_text(raw_text, filename, page_texts)
        return chunks

    @classmethod
    def _parse_pdf(cls, file_path: str, file_bytes: bytes = None) -> tuple:
        """Extract text from a PDF file."""
        if PdfReader is None:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")

        import io
        if file_bytes:
            reader = PdfReader(io.BytesIO(file_bytes))
        else:
            reader = PdfReader(file_path)

        page_texts = []
        full_text = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                page_texts.append({"page": i + 1, "text": text})
                full_text.append(text)

        return "\n\n".join(full_text), page_texts

    @classmethod
    def _parse_docx(cls, file_path: str, file_bytes: bytes = None) -> tuple:
        """Extract text from a DOCX file."""
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX parsing.")

        import io
        if file_bytes:
            doc = DocxDocument(io.BytesIO(file_bytes))
        else:
            doc = DocxDocument(file_path)

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        return full_text, [{"page": 1, "text": full_text}]

    @classmethod
    def _parse_text(cls, file_path: str, file_bytes: bytes = None) -> tuple:
        """Extract text from a plain text or markdown file."""
        if file_bytes:
            text = file_bytes.decode("utf-8", errors="replace")
        else:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()

        text = text.strip()
        return text, [{"page": 1, "text": text}]

    @classmethod
    def _chunk_text(cls, text: str, filename: str, page_texts: list) -> List[DocumentChunk]:
        """Split text into overlapping chunks."""
        if not text:
            return []

        # Clean up excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)

        chunks = []
        start = 0
        chunk_index = 0

        while start < len(text):
            end = start + cls.CHUNK_SIZE

            # Try to break at a sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                search_start = max(start + cls.CHUNK_SIZE - 100, start)
                search_end = min(start + cls.CHUNK_SIZE + 50, len(text))
                segment = text[search_start:search_end]

                # Find the best break point
                for pattern in ['. ', '.\n', '\n\n', '\n', '; ', ', ']:
                    pos = segment.rfind(pattern)
                    if pos != -1:
                        end = search_start + pos + len(pattern)
                        break

            chunk_text = text[start:end].strip()

            if chunk_text:
                # Determine page number
                page_num = cls._find_page_number(start, text, page_texts)

                chunks.append(DocumentChunk(
                    text=chunk_text,
                    chunk_index=chunk_index,
                    source_file=filename,
                    page_number=page_num,
                    start_char=start,
                    end_char=end,
                ))
                chunk_index += 1

            start = end - cls.CHUNK_OVERLAP
            if start >= len(text):
                break

        return chunks

    @classmethod
    def _find_page_number(cls, char_pos: int, full_text: str, page_texts: list) -> int:
        """Determine which page a character position belongs to."""
        if not page_texts:
            return 1

        cumulative = 0
        for pt in page_texts:
            cumulative += len(pt["text"]) + 2  # +2 for \n\n separator
            if char_pos < cumulative:
                return pt["page"]

        return page_texts[-1]["page"]
